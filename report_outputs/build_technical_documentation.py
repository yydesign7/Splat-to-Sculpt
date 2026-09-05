from __future__ import annotations

import re
import textwrap
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "Splat_to_Sculpt_Technical_Documentation.md"
OUTPUT = ROOT / "Splat_to_Sculpt_Technical_Documentation.docx"
ASSET_DIR = ROOT / "technical_documentation_assets"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "243447"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GREY = "F2F4F7"
GRID = "C8D2DE"
WHITE = "FFFFFF"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def pil_colour(value: str) -> str:
    return f"#{value}" if re.fullmatch(r"[0-9A-Fa-f]{6}", value) else value


def font_path(bold: bool = False) -> str:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Helvetica.ttc",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return candidate
    raise FileNotFoundError("No suitable diagram font found")


def rounded_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], fill: str, outline: str, radius: int = 18, width: int = 3) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=pil_colour(fill), outline=pil_colour(outline), width=width)


def centred_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font: ImageFont.FreeTypeFont, fill: str, max_chars: int = 22) -> None:
    lines = textwrap.wrap(text, width=max_chars) or [text]
    spacing = 7
    line_boxes = [draw.textbbox((0, 0), line, font=font) for line in lines]
    heights = [b[3] - b[1] for b in line_boxes]
    total_h = sum(heights) + spacing * (len(lines) - 1)
    y = box[1] + (box[3] - box[1] - total_h) / 2
    for line, b, h in zip(lines, line_boxes, heights):
        w = b[2] - b[0]
        x = box[0] + (box[2] - box[0] - w) / 2
        draw.text((x, y), line, font=font, fill=pil_colour(fill))
        y += h + spacing


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill: str = "52708F", width: int = 5) -> None:
    fill = pil_colour(fill)
    draw.line([start, end], fill=fill, width=width)
    dx, dy = end[0] - start[0], end[1] - start[1]
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    size = 14
    p1 = (end[0] - ux * size + px * size * 0.55, end[1] - uy * size + py * size * 0.55)
    p2 = (end[0] - ux * size - px * size * 0.55, end[1] - uy * size - py * size * 0.55)
    draw.polygon([end, p1, p2], fill=fill)


def make_architecture_diagram(path: Path) -> None:
    image = Image.new("RGB", (1800, 980), pil_colour(WHITE))
    draw = ImageDraw.Draw(image)
    title = ImageFont.truetype(font_path(True), 44)
    heading = ImageFont.truetype(font_path(True), 28)
    body = ImageFont.truetype(font_path(False), 24)
    draw.text((70, 45), "Local-first system architecture", font=title, fill=f"#{DARK_BLUE}")

    boxes = {
        "browser": (80, 180, 420, 420),
        "next": (560, 180, 920, 420),
        "tools": (1080, 110, 1710, 490),
        "session": (500, 610, 930, 850),
        "assets": (1110, 610, 1600, 850),
    }
    rounded_box(draw, boxes["browser"], "EAF1F8", BLUE)
    rounded_box(draw, boxes["next"], "EDF5F0", "4A7C65")
    rounded_box(draw, boxes["tools"], "F4F1F8", "775A9A")
    rounded_box(draw, boxes["session"], "F7F3E8", "9B7A2E")
    rounded_box(draw, boxes["assets"], "F8EEEE", "9A5962")

    centred_text(draw, (80, 190, 420, 265), "Browser workflow", heading, f"#{DARK_BLUE}")
    centred_text(draw, (95, 275, 405, 405), "React Flow canvas\nNode state + previews\nRun / Stop / Clear", body, "#243447")
    centred_text(draw, (560, 190, 920, 265), "Next.js local API", heading, "#315F4B")
    centred_text(draw, (575, 275, 905, 405), "Validation + routing\nTask stores + polling\nProcess management", body, "#243447")
    centred_text(draw, (1080, 120, 1710, 190), "Local processing tools", heading, "#60437F")

    tool_boxes = [(1120, 220, 1360, 315), (1430, 220, 1670, 315), (1120, 350, 1360, 445), (1430, 350, 1670, 445)]
    for box, label in zip(tool_boxes, ["FFmpeg + COLMAP", "Python + Nerfstudio", "Blender", "ComfyUI"]):
        rounded_box(draw, box, WHITE, "B6A6C8", radius=12, width=2)
        centred_text(draw, box, label, body, "#3C3150", max_chars=18)

    centred_text(draw, (500, 625, 930, 700), "Ephemeral session", heading, "#7A5A00")
    centred_text(draw, (520, 710, 910, 835), "Session work files\nFrames, masks, COLMAP,\nmesh layers and job outputs", body, "#403A2A", max_chars=30)
    centred_text(draw, (1110, 625, 1600, 700), "Published Assets", heading, "#7A3E48")
    centred_text(draw, (1130, 710, 1580, 835), "public/asset-published\nReusable videos, splats,\nmodels and final renders", body, "#4A3034")

    arrow(draw, (420, 300), (560, 300))
    arrow(draw, (920, 300), (1080, 300))
    arrow(draw, (740, 420), (715, 610))
    arrow(draw, (930, 730), (1110, 730))
    arrow(draw, (1395, 490), (1370, 610))
    image.save(path, quality=95)


def make_lifecycle_diagram(path: Path) -> None:
    image = Image.new("RGB", (1800, 760), pil_colour(WHITE))
    draw = ImageDraw.Draw(image)
    title = ImageFont.truetype(font_path(True), 44)
    heading = ImageFont.truetype(font_path(True), 27)
    body = ImageFont.truetype(font_path(False), 22)
    draw.text((70, 45), "Runtime file lifecycle", font=title, fill=f"#{DARK_BLUE}")

    steps = [
        ((70, 230, 345, 455), "1  Create session", "UUID identifies one\nworkflow runtime"),
        ((430, 230, 735, 455), "2  Write temporary files", "Uploads, frames, masks,\nCOLMAP, meshes, previews"),
        ((820, 230, 1115, 455), "3  Publish selected output", "Copy reusable result to\nasset-published"),
        ((1200, 230, 1510, 455), "4  Register metadata", "Add one AssetEntry to\nassets.json"),
    ]
    fills = ["EAF1F8", "F7F3E8", "EDF5F0", "F8EEEE"]
    outlines = [BLUE, "9B7A2E", "4A7C65", "9A5962"]
    for (box, head, desc), fill, outline in zip(steps, fills, outlines):
        rounded_box(draw, box, fill, outline)
        centred_text(draw, (box[0] + 10, box[1] + 20, box[2] - 10, box[1] + 90), head, heading, f"#{outline}", 24)
        centred_text(draw, (box[0] + 15, box[1] + 105, box[2] - 15, box[3] - 20), desc, body, "#273444", 25)
    for a, b in zip(steps, steps[1:]):
        arrow(draw, (a[0][2] + 15, 342), (b[0][0] - 15, 342))

    rounded_box(draw, (440, 555, 1090, 690), "F2F4F7", "8793A1", radius=15, width=2)
    centred_text(draw, (455, 565, 1075, 680), "Clear or normal exit removes session files.\nPublished Assets remain available.", heading, "#435163", 44)
    arrow(draw, (585, 455), (650, 555), fill="8793A1", width=4)
    image.save(path, quality=95)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)


def set_run_font(run, name: str = "Calibri", size: float | None = None, colour: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if colour is not None:
        run.font.color.rgb = RGBColor.from_string(colour)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_inline(paragraph, text: str, size: float | None = None, colour: str | None = None) -> None:
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`)")
    for part in pattern.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, colour=colour, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Consolas", size=(size or 11) - 0.5, colour=DARK_BLUE)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "F1F4F7")
            run._element.get_or_add_rPr().append(shd)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, colour=colour)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_end])


def add_toc(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("Contents")
    set_run_font(run, size=20, colour=DARK_BLUE, bold=True)

    entries = [
        ("Document Conventions", 3),
        ("1. Overview", 3),
        ("2. System Architecture", 6),
        ("3. Workflow Engine and Data Contracts", 9),
        ("4. Reconstruction Pipeline", 11),
        ("5. Mesh and Model Processing", 16),
        ("6. Local Tool Integration", 19),
        ("7. Storage and Asset Lifecycle", 21),
        ("8. API Reference", 23),
        ("9. Frontend Implementation", 27),
        ("10. Installation and Deployment", 29),
        ("11. Testing and Verification", 32),
        ("12. Troubleshooting and Maintenance", 34),
        ("Appendices", 38),
    ]
    for label, page in entries:
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.35), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )
        run = p.add_run(f"{label}\t{page}")
        set_run_font(
            run,
            size=11,
            colour=INK,
            bold=bool(re.match(r"^(?:\d+\.|Appendices)", label)),
        )


def set_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, colour, before, after) in specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(colour)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_cover(document: Document) -> None:
    for _ in range(5):
        document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("DEVELOPMENT HANDOVER")
    set_run_font(r, size=11, colour=BLUE, bold=True)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Splat to Sculpt")
    set_run_font(r, size=30, colour=DARK_BLUE, bold=True)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(32)
    r = p.add_run("Technical Documentation")
    set_run_font(r, size=17, colour=BLUE)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("A local-first node workflow for video reconstruction, mesh processing and generative presentation")
    set_run_font(r, size=11, colour=MUTED, italic=True)
    for _ in range(5):
        document.add_paragraph()
    for label, value in (
        ("Prepared by", "Yu Yi"),
        ("Document purpose", "Development handover and maintenance reference"),
        ("Version", "1.0.0"),
        ("Version basis", "GitHub Release v1.0.0, commit 5f6e35f, 3 August 2026"),
    ):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{label}: ")
        set_run_font(r, size=10.5, colour=MUTED, bold=True)
        r = p.add_run(value)
        set_run_font(r, size=10.5, colour=INK)
    document.add_page_break()


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def table_widths(rows: list[list[str]]) -> list[int]:
    columns = len(rows[0])
    scores = []
    for index in range(columns):
        values = [len(re.sub(r"[`*]", "", row[index])) for row in rows if index < len(row)]
        scores.append(max(8, min(34, max(values, default=10))))
    minimum = 780 if columns >= 6 else 900
    widths = [max(minimum, round(CONTENT_DXA * score / sum(scores))) for score in scores]
    diff = CONTENT_DXA - sum(widths)
    widths[-1] += diff
    if widths[-1] < minimum:
        deficit = minimum - widths[-1]
        widths[-1] = minimum
        for i in sorted(range(columns - 1), key=lambda x: widths[x], reverse=True):
            take = min(deficit, max(0, widths[i] - minimum))
            widths[i] -= take
            deficit -= take
            if deficit <= 0:
                break
    widths[-1] += CONTENT_DXA - sum(widths)
    return widths


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    widths = table_widths(rows)
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    font_size = 7.8 if len(widths) >= 6 else 8.5 if len(widths) >= 4 else 9.0

    for row_index, values in enumerate(rows):
        for col_index, value in enumerate(values):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                shade_cell(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            if len(value) < 18 and col_index > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(p, value, size=font_size, colour=INK)
            if row_index == 0:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_code_block(document: Document, code: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.right_indent = Inches(0.14)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GREY)
    p_pr.append(shd)
    run = p.add_run(code.rstrip())
    set_run_font(run, name="Consolas", size=8.4, colour="26384A")


def add_figure(document: Document, key: str, figure_number: int, paths: dict[str, Path]) -> None:
    captions = {
        "architecture": "Local-first architecture and processing boundaries.",
        "lifecycle": "Temporary-session and published-asset lifecycle.",
    }
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    inline_shape = run.add_picture(str(paths[key]), width=Inches(6.35))
    inline_shape._inline.docPr.set("title", captions[key])
    inline_shape._inline.docPr.set("descr", captions[key])
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(f"Figure {figure_number}. {captions[key]}")
    set_run_font(r, size=9.5, colour=MUTED, italic=True)


def build_document() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    architecture = ASSET_DIR / "architecture.png"
    lifecycle = ASSET_DIR / "file_lifecycle.png"
    make_architecture_diagram(architecture)
    make_lifecycle_diagram(lifecycle)
    figure_paths = {"architecture": architecture, "lifecycle": lifecycle}

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True
    set_styles(document)

    document.core_properties.title = "Splat to Sculpt Technical Documentation"
    document.core_properties.subject = "Development handover and maintenance reference"
    document.core_properties.author = "Yu Yi"
    document.core_properties.keywords = "Splat to Sculpt, 3DGS, COLMAP, Blender, ComfyUI, technical documentation"

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("SPLAT TO SCULPT  /  TECHNICAL DOCUMENTATION")
    set_run_font(hr, size=8.5, colour=MUTED, bold=True)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run("Development handover  |  ")
    set_run_font(fr, size=8.5, colour=MUTED)
    add_page_field(fp)

    settings = document.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    add_cover(document)
    add_toc(document)
    document.add_page_break()

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    start = next(i for i, line in enumerate(lines) if line.strip() == "## Document Conventions")
    lines = lines[start:]
    i = 0
    in_code = False
    code_lines: list[str] = []
    figure_number = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                add_code_block(document, "\n".join(code_lines))
                in_code = False
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if not stripped or stripped == "---":
            i += 1
            continue
        if stripped.startswith("[[FIGURE:"):
            key = stripped.removeprefix("[[FIGURE:").removesuffix("]]" )
            figure_number += 1
            add_figure(document, key, figure_number, figure_paths)
            i += 1
            continue
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\s*\|(?:\s*:?-+:?\s*\|)+\s*$", lines[i + 1]):
            table_lines = [stripped]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            add_table(document, parse_table(table_lines))
            continue
        if stripped.startswith("### "):
            document.add_paragraph(stripped[4:], style="Heading 3")
            i += 1
            continue
        if stripped.startswith("## "):
            text = stripped[3:]
            style = "Heading 1" if text == "Document Conventions" else "Heading 2"
            document.add_paragraph(text, style=style)
            i += 1
            continue
        if stripped.startswith("# "):
            p = document.add_paragraph(stripped[2:], style="Heading 1")
            if stripped.startswith("# Appendices"):
                p.paragraph_format.page_break_before = True
            i += 1
            continue
        if re.match(r"^-\s+", stripped):
            p = document.add_paragraph(style="List Bullet")
            add_inline(p, re.sub(r"^-\s+", "", stripped))
            i += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            p = document.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        p = document.add_paragraph()
        add_inline(p, stripped)
        i += 1

    document.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
