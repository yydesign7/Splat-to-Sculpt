from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import re
import sys

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Pt


INPUT_PATH = Path(
    "/Users/yuyi/Desktop/"
    "Splat_to_Sculpt_Capstone_Final_Report_Bilingual_Complete_Translation_Illustrated.docx"
)
OUTPUT_PATH = Path(
    "/Users/yuyi/projects/report_output/"
    "Splat_to_Sculpt_Capstone_Final_Report_Bilingual_UPDATED_ComfyUI.docx"
)
FIGURES_DIR = Path("/Users/yuyi/projects/report_assets/final_figures_current")

EXPECTED_TOC = [
    "Abstract",
    "Acknowledgement",
    "Statement of Contributions",
    "List of Figures",
    "List of Tables",
    "1. Introduction",
    "2. Background and Related Work",
    "3. Design Requirements and System Concept",
    "4. Workflow Design",
    "5. Website Function Design",
    "6. UI Design",
    "7. Implementation",
    "8. Results and Discussion",
    "9. Conclusions",
    "10. Future Work",
    "References",
    "Appendices",
]


REPLACEMENTS: dict[int, str] = {
    8: (
        "This report presents the design and implementation of Splat to Sculpt, a node-based "
        "web application that converts captured video or an existing PLY source into Gaussian "
        "Splat assets, conventional mesh models, surface-processed outputs, and generated videos. "
        "The project addresses the gap between advanced reconstruction tools and creative users "
        "who need a visible, controllable, and reusable production workflow. The current system "
        "integrates video upload, frame extraction, COLMAP camera reconstruction, capability-aware "
        "Gaussian generation, foreground masks, multi-view point filtering, splat preview, mesh "
        "conversion, model cleanup, surface processing, asset management, workflow saving, local "
        "ComfyUI video generation, and final video preview. CUDA systems can target Nerfstudio "
        "Splatfacto True training, while Apple MPS and CPU systems use a clearly identified Fast "
        "Initializer route. Direct PLY input remains on the initializer route because it lacks the "
        "images and recovered camera poses required for multi-view training. A new ComfyUI Video "
        "Gen node receives the processed model, checks the local ComfyUI and Seedance environment, "
        "submits an API workflow, and returns the generated video to the final preview node. The "
        "interface exposes device, PLY target, progress, dependency status, and recovery actions "
        "instead of hiding technically different processes behind one generate command. Together "
        "with on-demand Three.js rendering and a reusable Figma UI Kit, these features form a "
        "design-led integration layer that connects technical truth, recoverability, workflow "
        "semantics, and visual consistency."
    ),
    9: (
        "中文对照：本报告介绍 Splat to Sculpt 的设计与实现。该项目是一款节点式网页应用，"
        "可把拍摄视频或已有 PLY 数据源转换为 Gaussian Splat 资产、传统 mesh 模型、经过表面"
        "处理的输出以及生成视频。项目所回应的问题，是先进三维重建工具与创意用户实际需求之间"
        "的差距：用户需要一套可见、可控制且可复用的制作工作流。当前系统整合了视频上传、帧提取、"
        "COLMAP 相机重建、具备设备能力判断的 Gaussian 生成、前景遮罩、多视角点过滤、splat "
        "预览、mesh 转换、模型清理、表面处理、资产管理、工作流保存、本地 ComfyUI 视频生成和"
        "最终视频预览。在 CUDA 系统上，工作流可以使用 Nerfstudio Splatfacto 进行 True "
        "training；在 Apple MPS 和 CPU 系统上，则使用被清楚标识的 Fast Initializer 路径。"
        "直接输入 PLY 时，由于缺少多视角训练所需的图像和恢复出的相机位姿，系统会保持在 "
        "initializer 路径。新增的 ComfyUI Video Gen 节点接收经过处理的模型，检查本地 "
        "ComfyUI 与 Seedance 环境，通过 API 提交工作流，并把生成视频传递到最终预览节点。"
        "界面会公开设备、PLY target、进度、依赖状态和恢复操作，而不是把技术上不同的过程隐藏"
        "在同一个生成命令之后。结合按需 Three.js 渲染和可复用 Figma UI Kit，这些功能构成了"
        "一个以设计为主导的整合层，把技术真实性、可恢复性、工作流语义和视觉一致性连接起来。"
    ),
    21: (
        "• Added capability-aware True training and Fast Initializer routes, protected direct-PLY "
        "input from unsupported training, and integrated local ComfyUI/Seedance model-to-video output."
    ),
    22: (
        "中文对照：• 增加了具备设备能力判断的 True training 与 Fast Initializer 路径，"
        "防止直接 PLY 进入不受支持的训练流程，并接入本地 ComfyUI/Seedance 的模型到视频输出。"
    ),
    80: (
        "The objectives are: first, to define a coherent end-to-end workflow from video or PLY input "
        "to Gaussian Splat, mesh, surface, and generated-video outputs; second, to implement node "
        "components for upload, frame extraction, Gaussian generation, mesh generation, model cleanup, "
        "surface processing, ComfyUI video generation, and video preview; third, to design port and "
        "file-type semantics so that point clouds, splats, meshes, GLB models, and videos are not "
        "confused; fourth, to support long-running computation with truthful progress and stop/recovery "
        "behaviour; fifth, to expose local ComfyUI connection and dependency state without overwhelming "
        "the main canvas; sixth, to maintain visual consistency through a Figma component system; and "
        "seventh, to document limitations honestly, especially the difference between initializer output "
        "and CUDA Splatfacto training and the dependency of video generation on a configured local service."
    ),
    81: (
        "中文对照：项目目标包括：第一，定义从视频或 PLY 输入到 Gaussian Splat、mesh、表面"
        "处理结果和生成视频的一致端到端工作流；第二，实现上传、帧提取、Gaussian 生成、mesh "
        "生成、模型清理、表面处理、ComfyUI 视频生成和视频预览等节点组件；第三，设计端口与文件"
        "类型语义，避免混淆点云、splats、meshes、GLB 模型和视频；第四，通过符合事实的进度与"
        "停止/恢复行为支持长时间运算；第五，在不使主画布过载的前提下公开本地 ComfyUI 连接和"
        "依赖状态；第六，通过 Figma 组件系统保持视觉一致性；第七，如实记录系统限制，特别是 "
        "initializer 输出与 CUDA Splatfacto 训练之间的差异，以及视频生成对配置正确的本地"
        "服务的依赖。"
    ),
    83: (
        "The deliverable is a working browser-based prototype rather than a standalone desktop "
        "application. The interface runs through a local development server and uses Next.js API "
        "routes to coordinate Python, COLMAP, Nerfstudio, Open3D, Blender, and a local ComfyUI "
        "instance. Within the intended scope, the application supports video upload, configurable "
        "frame extraction, foreground-aware COLMAP reconstruction, Gaussian Splat generation, direct "
        "PLY initialization, splat preview, splat-to-mesh conversion, model cleanup, surface and "
        "lighting adjustment, asset saving with thumbnails, workflow saving, non-destructive Clear, "
        "ComfyUI/Seedance model-to-video generation, and final video preview. The protected Default "
        "Workflow connects these responsibilities in one reusable graph, while Material Gen remains "
        "registered for compatibility but is excluded from the preset while its external service is disabled."
    ),
    84: (
        "中文对照：项目交付物是一个可运行的浏览器原型，而不是独立桌面应用。界面通过本地开发"
        "服务器运行，并使用 Next.js API 路由协调 Python、COLMAP、Nerfstudio、Open3D、"
        "Blender 和本地 ComfyUI。预定范围内，应用支持视频上传、可配置帧提取、具备前景感知的 "
        "COLMAP 重建、Gaussian Splat 生成、直接 PLY 初始化、splat 预览、splat-to-mesh "
        "转换、模型清理、表面与灯光调整、带缩略图的资产保存、工作流保存、非破坏性 Clear、"
        "ComfyUI/Seedance 模型到视频生成以及最终视频预览。受保护的 Default Workflow 把这些"
        "职责连接成一张可复用节点图；Material Gen 仍为兼容旧工作流而保留注册，但在其外部服务"
        "停用期间不会进入默认预设。"
    ),
    85: (
        "The current implementation remains a local browser prototype rather than a hosted production "
        "service. CUDA True training depends on Nerfstudio, gsplat, and configured executables; MPS/CPU "
        "output remains an initializer representation. Direct PLY input does not provide camera poses, "
        "foreground masks have not been quantitatively benchmarked, mesh conversion can amplify defects "
        "in sparse point data, and the browser splat preview is approximate. ComfyUI Video Gen additionally "
        "requires a running local ComfyUI server, the expected Seedance custom nodes and workflows, and "
        "available model or API capability. Stopping the website workflow does not yet guarantee interruption "
        "of a prompt that ComfyUI has already accepted into its own queue. These boundaries are shown as "
        "product limitations rather than hidden behind a generic success state."
    ),
    86: (
        "中文对照：当前实现仍是本地浏览器原型，而不是已托管的生产服务。CUDA True training "
        "依赖 Nerfstudio、gsplat 以及配置正确的可执行文件；MPS/CPU 输出仍属于 initializer "
        "表示。直接 PLY 不提供相机位姿，前景遮罩尚未经过定量基准测试，mesh 转换可能放大稀疏"
        "点数据中的缺陷，浏览器 splat 预览也只是近似效果。ComfyUI Video Gen 还要求本地 "
        "ComfyUI 服务正在运行、预期的 Seedance custom nodes 与 workflows 已经安装，并且"
        "所需模型或 API 能力可用。停止网页工作流目前不能保证中断已经被 ComfyUI 接收到其内部"
        "队列中的 prompt。报告把这些边界明确作为产品限制，而不是隐藏在通用成功状态之后。"
    ),
    87: (
        "Fig. 1.1. Current editor with the connected Default Workflow and ComfyUI Video Gen."
    ),
    88: (
        "Fig. 1.1 中文图注：当前编辑器展示已连接的 Default Workflow 与 ComfyUI Video Gen。"
    ),
    121: (
        "The main use scenario begins with a user recording a short mobile video of an object, uploading "
        "it into Video Upload, selecting a frame count, and running the Default Workflow. The system "
        "extracts frames, reconstructs cameras, generates a splat, converts it to a mesh, cleans the model, "
        "and applies surface or lighting changes. Surface Processing then supplies the final model to "
        "ComfyUI Video Gen, where the user can inspect the local connection, synchronize a Seedance preset, "
        "adjust prompt and video parameters, and generate a presentation video that appears in Video Preview. "
        "A secondary scenario begins from a reusable PLY or model asset, while a third scenario saves and "
        "reopens an edited node graph for later experimentation."
    ),
    122: (
        "中文对照：主要使用场景从用户用手机录制物体短视频开始。用户把视频上传到 Video Upload，"
        "选择帧数，然后运行 Default Workflow。系统会提取帧、重建相机、生成 splat、把它转换"
        "为 mesh、清理模型，并应用表面或灯光修改。随后，Surface Processing 把最终模型提供给 "
        "ComfyUI Video Gen；用户可以在该节点中检查本地连接、同步 Seedance preset、调整 "
        "prompt 和视频参数，并生成在 Video Preview 中显示的展示视频。第二种场景从可复用的 "
        "PLY 或模型资产开始，第三种场景则保存并重新打开修改后的节点图，以便之后继续实验。"
    ),
    124: (
        "The functional requirements include two Gaussian input routes, foreground-aware frame "
        "reconstruction, native splat preview, splat-to-mesh conversion, model cleanup, surface and lighting "
        "controls, local ComfyUI video generation, final video preview, task cancellation, global Clear, "
        "asset thumbnails, and workflow presets. Direct PLY input must not enter True training without images "
        "and cameras. ComfyUI Video Gen must accept a supported model file, expose connection and Seedance "
        "readiness, submit a configurable API workflow, and return a reusable video URL. Material Gen must "
        "fail explicitly while its service is disabled and must not be required by the Default Workflow."
    ),
    125: (
        "中文对照：功能需求包括两条 Gaussian 输入路径、具备前景感知的帧重建、原生 splat "
        "预览、splat-to-mesh 转换、模型清理、表面与灯光控制、本地 ComfyUI 视频生成、最终"
        "视频预览、任务取消、全局 Clear、资产缩略图和工作流预设。缺少图像和相机数据时，直接 "
        "PLY 不得进入 True training。ComfyUI Video Gen 必须接收受支持的模型文件，公开连接"
        "与 Seedance 准备状态，提交可配置的 API workflow，并返回可复用视频 URL。当 "
        "Material Gen 服务停用时，该节点必须明确失败，而且 Default Workflow 不得依赖它。"
    ),
    126: (
        "Controls must correspond to active backend behaviour. Device and PLY target explain the Gaussian "
        "compute route, and training progress uses parsed iterations only when the backend supplies meaningful "
        "values. ComfyUI controls expose the server URL, detected input and output folders, Seedance pack "
        "status, prompt, resolution, ratio, duration, audio, seed, watermark, and render settings because these "
        "values affect the submitted workflow. Model URLs and lighting parameters must survive the handoff from "
        "Surface Processing, while videoUrl must be forwarded only through the connected video-output edge."
    ),
    127: (
        "中文对照：控件必须与真实后端行为对应。Device 和 PLY target 用于解释 Gaussian 计算"
        "路径，只有后端提供有意义的数值时，训练进度才使用解析出的迭代次数。ComfyUI 控件会公开"
        "服务器 URL、检测到的输入和输出文件夹、Seedance pack 状态、prompt、分辨率、比例、"
        "时长、音频、seed、watermark 和渲染设置，因为这些数值会影响提交的 workflow。"
        "model URL 与灯光参数必须在 Surface Processing 的交接中保留下来，而 videoUrl 只能"
        "通过已经连接的 video-output edge 向下游传递。"
    ),
    129: (
        "Non-functional requirements are transparency, recoverability, cross-platform tolerance, responsive "
        "interaction, dependency visibility, and bounded idle resource use. Long-running reconstruction must "
        "expose real stages and terminate tracked local processes after Stop. Clear must remove stale previews "
        "and runtime file references. ComfyUI connection failures, missing custom nodes, unavailable workflows, "
        "unsupported model formats, and queue limitations must be explained before or at the point of failure. "
        "WebGL previews should render when content or interaction changes rather than continuously consuming "
        "GPU resources in an idle browser."
    ),
    130: (
        "中文对照：非功能需求包括透明性、可恢复性、跨平台容错、响应式交互、依赖可见性以及受控"
        "的空闲资源占用。长时间重建必须显示真实阶段，并在 Stop 后终止被追踪的本地进程。Clear "
        "必须移除残留预览和运行时文件引用。ComfyUI 连接失败、缺少 custom nodes、workflow "
        "不可用、模型格式不受支持以及队列限制，都必须在失败前或失败发生处得到解释。WebGL 预览"
        "应只在内容或交互发生变化时渲染，而不是在浏览器空闲时持续消耗 GPU 资源。"
    ),
    134: (
        "The concept can be summarized as a node-based 3D production desk. The canvas is the main working "
        "area, the sidebar supplies nodes, assets, and saved workflows, and the top bar controls global "
        "Run/Stop, Save Workflow, and Clear actions. Each node is a small tool with typed handles, a preview "
        "surface, status feedback, and only the parameters owned by that stage. Edges encode both data movement "
        "and execution dependency. The output stage is represented explicitly: Surface Processing produces the "
        "model consumed by ComfyUI Video Gen, and the generated video travels through a video-specific edge to "
        "Video Preview."
    ),
    135: (
        "中文对照：系统概念可以概括为一张节点式三维制作台。画布是主要工作区，侧栏提供节点、"
        "资产和已保存工作流，顶部栏控制全局 Run/Stop、Save Workflow 和 Clear。每个节点都是"
        "一个小工具，具有带类型的 handles、预览表面、状态反馈以及只属于该阶段的参数。Edges "
        "同时编码数据移动和执行依赖。输出阶段被明确表示：Surface Processing 生成由 ComfyUI "
        "Video Gen 接收的模型，生成视频再通过视频专用 edge 传递到 Video Preview。"
    ),
    143: (
        "The current Default Workflow is arranged as a left-to-right production line: Video Upload, Frame "
        "Extraction, Gaussian Splat Gen, Mesh Gen, Model Cleanup, Surface Processing, ComfyUI Video Gen, and "
        "Video Preview. Sticky Notes remain outside the computational chain and explain graph use. The preset "
        "is available in the Workflows Library as a readonly baseline, allowing users to restore the intended "
        "topology without deleting their own saved workflows. Material Gen is excluded because its external "
        "generation service is disabled; local material and lighting controls remain available in Surface "
        "Processing before the model is passed to ComfyUI."
    ),
    144: (
        "中文对照：当前 Default Workflow 被组织为一条从左到右的制作线：Video Upload、Frame "
        "Extraction、Gaussian Splat Gen、Mesh Gen、Model Cleanup、Surface Processing、"
        "ComfyUI Video Gen 和 Video Preview。Sticky Notes 位于计算链之外，用于解释节点图"
        "使用方式。该预设作为只读基线保存在 Workflows Library 中，使用户可以恢复预期拓扑，"
        "而不会删除自己保存的工作流。Material Gen 因外部生成服务停用而被排除；模型传给 "
        "ComfyUI 之前，本地材质和灯光控制仍可在 Surface Processing 中使用。"
    ),
    145: (
        "Two representation boundaries are particularly important. Gaussian Splat Gen connects from its "
        "splat-output handle to Mesh Gen because a splat PLY and a conventional mesh are not interchangeable "
        "even when both use the PLY container. Mesh Gen creates a GLB, OBJ, or mesh PLY that can be processed "
        "by conventional model tools. The second boundary appears between Surface Processing and ComfyUI "
        "Video Gen: a model edge carries the final supported model, while the ComfyUI node returns a video "
        "rather than another editable model. Typed handles make both transformations visible."
    ),
    146: (
        "中文对照：其中有两处表示边界尤其重要。Gaussian Splat Gen 从 splat-output handle "
        "连接到 Mesh Gen，因为即使 splat PLY 和传统 mesh 都可能使用 PLY 容器，二者也不能"
        "互换。Mesh Gen 会创建可由传统模型工具处理的 GLB、OBJ 或 mesh PLY。第二处边界位于 "
        "Surface Processing 与 ComfyUI Video Gen 之间：model edge 携带最终且受支持的模型，"
        "ComfyUI 节点返回的是视频，而不是另一份可编辑模型。带类型的 handles 让两次转换都清楚"
        "可见。"
    ),
    153: (
        "Run executes the connected graph rather than broadcasting data to every node on the canvas. The "
        "workflow engine maps source handles to output fields, maps target handles to receiving fields, checks "
        "required inputs, and pushes data only across existing edges. ComfyUI Video Gen therefore starts only "
        "after it receives a valid final model from an upstream connection, and Video Preview receives only "
        "the returned videoUrl. Terminal nodes are monitored so the global workflow returns to idle when every "
        "connected endpoint has finished or reported an error."
    ),
    154: (
        "中文对照：Run 只执行已连接的节点图，而不会把数据广播到画布中的每个节点。工作流引擎"
        "把 source handles 映射到输出字段，把 target handles 映射到接收字段，检查必需输入，"
        "并且只沿现有 edges 推送数据。因此，ComfyUI Video Gen 只有在通过上游连线接收到有效"
        "最终模型后才会开始，而 Video Preview 只接收返回的 videoUrl。系统会监测末端节点，"
        "所以当所有已连接终点完成或报告错误时，全局工作流会恢复到 idle。"
    ),
    155: (
        "Stop and Clear operate at different levels. Stop changes global execution state and calls cancellation "
        "routes for tracked COLMAP, Python, Blender, and Nerfstudio processes. Clear resets uploaded files, "
        "generated URLs, errors, progress, previews, and task identifiers while preserving node positions, "
        "edges, saved topology, and annotations. The ComfyUI integration has a narrower cancellation boundary: "
        "the website can stop waiting and prevent downstream delivery, but a prompt already accepted by the "
        "separate ComfyUI queue may continue until ComfyUI itself is interrupted."
    ),
    156: (
        "中文对照：Stop 和 Clear 作用于不同层级。Stop 会修改全局执行状态，并调用取消路由终止"
        "被追踪的 COLMAP、Python、Blender 和 Nerfstudio 进程。Clear 会重置上传文件、生成 "
        "URL、错误、进度、预览和任务标识符，同时保留节点位置、edges、已保存拓扑和注释。"
        "ComfyUI 接入具有更窄的取消边界：网站可以停止等待并阻止结果继续向下游传递，但已经被"
        "独立 ComfyUI 队列接收的 prompt 可能继续运行，直到在 ComfyUI 中另行中断。"
    ),
    162: "4.5 Downstream Data and Model-to-Video Handoff",
    163: (
        "Gaussian Splat Gen now focuses on reconstruction output. It returns a splatUrl, sourcePlyUrl, "
        "device and backend metadata, targetPlyType, progress, and error state. Mesh Gen receives the connected "
        "splat or point source, creates conventional geometry, and returns modelUrl, output type, face count, "
        "render information, and any mesh-region outputs produced after reconstruction. Model Cleanup and "
        "Surface Processing preserve the selected model URL and lighting data instead of reintroducing "
        "Gaussian-specific processing."
    ),
    164: (
        "中文对照：Gaussian Splat Gen 现在专注于重建输出。它会返回 splatUrl、sourcePlyUrl、"
        "设备与后端 metadata、targetPlyType、进度和错误状态。Mesh Gen 接收通过连线传入的 "
        "splat 或点数据源，创建传统几何，并返回 modelUrl、输出类型、面数、渲染信息，以及在"
        "重建之后生成的 mesh 区域输出。Model Cleanup 与 Surface Processing 会保留所选模型 "
        "URL 和灯光数据，而不会重新引入只属于 Gaussian 的处理。"
    ),
    165: (
        "The final handoff begins when Surface Processing exposes its processed model through obj-output. "
        "ComfyUI Video Gen accepts GLB, GLTF, OBJ, FBX, or BLEND input, stores the model reference and light "
        "context, and exposes settings that affect the local generation request. Its output is a videoUrl with "
        "a video name and prompt identifier. The workflow engine forwards this result through video-output to "
        "Video Preview, and the asset recorder can retain the generated clip as a reusable render-video asset."
    ),
    166: (
        "中文对照：最终交接从 Surface Processing 通过 obj-output 输出处理后的模型开始。"
        "ComfyUI Video Gen 接受 GLB、GLTF、OBJ、FBX 或 BLEND 输入，保存模型引用与灯光上下文，"
        "并公开会影响本地生成请求的设置。其输出包括 videoUrl、视频名称和 prompt identifier。"
        "工作流引擎会通过 video-output 把该结果转发到 Video Preview，资产记录器也可以把生成"
        "片段保存为可复用的 render-video 资产。"
    ),
    167: (
        "Mesh-region metadata remains a downstream geometry concern. When Mesh Gen exports "
        "geometry_graph_surface regions after triangle reconstruction, those files describe conventional "
        "mesh surfaces and can be forwarded for surface operations. They are not embedded layer identifiers "
        "inside the Gaussian PLY and do not change the meaning of Fast Initializer or True training. Keeping "
        "this distinction in the data contract prevents the interface from overstating what the splat contains."
    ),
    168: (
        "中文对照：Mesh 区域 metadata 仍属于下游几何处理。当 Mesh Gen 在三角网格重建之后导出 "
        "geometry_graph_surface 区域时，这些文件描述的是传统 mesh 表面，并可继续传给表面"
        "处理。它们不是嵌入 Gaussian PLY 内部的层级标识，也不会改变 Fast Initializer 或 True "
        "training 的含义。在数据契约中保持这一区分，可以防止界面夸大 splat 文件实际包含的内容。"
    ),
    170: (
        "Fig. 4.1. Current Default Workflow from video input through Gaussian generation, mesh processing, "
        "Surface Processing, ComfyUI Video Gen, and Video Preview."
    ),
    171: (
        "Fig. 4.1 中文图注：当前 Default Workflow，从视频输入经过 Gaussian 生成、mesh 处理和 "
        "Surface Processing，再到 ComfyUI Video Gen 与 Video Preview。"
    ),
    172: (
        "Fig. 4.2. Placeholder: Insert a data-flow diagram comparing the frame/COLMAP route with the "
        "direct-PLY Fast Initializer route and showing the shared model-to-ComfyUI video handoff."
    ),
    173: (
        "Fig. 4.2 需要展示的图片：对比 frames/COLMAP 路径与直接 PLY Fast Initializer 路径，"
        "并展示二者共用的模型到 ComfyUI 视频交接。"
    ),
    181: (
        "Each node combines typed ports, a stable preview box, controls relevant to its stage, and status "
        "feedback. Video Upload owns file selection and frame count; Frame Extraction owns image sampling; "
        "Gaussian Splat Gen owns device-aware reconstruction, PLY target, training mode, training budget, and "
        "splat preview; Mesh Gen owns representation conversion; Model Cleanup and Surface Processing own "
        "conventional model operations. ComfyUI Video Gen owns local server state, Seedance readiness, prompt "
        "and video settings, model submission, and generated-video recovery. Video Preview owns final playback "
        "and download. This division keeps unrelated technical parameters out of neighbouring stages."
    ),
    182: (
        "中文对照：每个节点都组合了带类型的端口、稳定预览框、与当前阶段相关的控件以及状态反馈。"
        "Video Upload 负责文件选择和帧数；Frame Extraction 负责图像采样；Gaussian Splat Gen "
        "负责设备感知重建、PLY target、训练模式、训练预算和 splat 预览；Mesh Gen 负责表示"
        "转换；Model Cleanup 与 Surface Processing 负责传统模型操作。ComfyUI Video Gen "
        "负责本地服务器状态、Seedance 准备情况、prompt 与视频设置、模型提交和生成视频回收；"
        "Video Preview 负责最终播放与下载。这样的分工使无关技术参数不会进入相邻阶段。"
    ),
    186: (
        "Next.js API routes form an integration layer between React nodes and external tools. Reconstruction "
        "routes create task identifiers and polling endpoints, while the ComfyUI route coordinates a separate "
        "local HTTP service. It validates the session and model, detects ComfyUI folders, copies the model to "
        "input/3d, builds a prompt from the synchronized preset, submits it to /prompt, polls /history by "
        "prompt identifier, retrieves the completed media through /view, and copies the video into the current "
        "ephemeral session. This design converts ComfyUI's file and queue conventions into the same URL-based "
        "contract used by the rest of the website."
    ),
    187: (
        "中文对照：Next.js API 路由在 React 节点与外部工具之间构成整合层。重建路由会创建 task "
        "identifiers 和轮询端点，而 ComfyUI 路由负责协调独立的本地 HTTP 服务。它会验证会话"
        "与模型、检测 ComfyUI 文件夹、把模型复制到 input/3d、根据已同步 preset 构建 prompt、"
        "提交到 /prompt、根据 prompt identifier 轮询 /history、通过 /view 获取完成媒体，"
        "再把视频复制到当前 ephemeral session。这样的设计把 ComfyUI 的文件和队列约定转换成"
        "网站其他部分使用的同一种 URL 数据契约。"
    ),
    188: (
        "Ephemeral session roots keep extracted frames, masks, COLMAP workspaces, splats, meshes, thumbnails, "
        "and copied ComfyUI videos separated by browser session. Published assets are recorded separately for "
        "reuse. Clear resets node references to runtime files and previews, including model-processing and "
        "video-generation nodes that could otherwise retain stale GLB or MP4 content, while the saved workflow "
        "topology remains intact. The source file placed in ComfyUI input/3d and any queued ComfyUI execution "
        "belong to the external service boundary and require separate lifecycle handling."
    ),
    189: (
        "中文对照：Ephemeral session roots 按浏览器会话隔离提取帧、遮罩、COLMAP 工作区、splats、"
        "meshes、缩略图以及复制回来的 ComfyUI 视频。已发布资产会单独记录，以便复用。Clear 会"
        "重置节点对运行时文件和预览的引用，包括可能保留过期 GLB 或 MP4 内容的模型处理与视频生成"
        "节点，同时保持已保存工作流拓扑不变。被放入 ComfyUI input/3d 的源文件和任何已进入队列"
        "的 ComfyUI 执行都属于外部服务边界，需要单独处理其生命周期。"
    ),
    191: (
        "The website exposes technical facts only when they affect user decisions. Device and PLY target "
        "explain whether a Gaussian run performs CUDA True training or initializer conversion, and current "
        "training step appears only when the backend supplies a meaningful count. ComfyUI Video Gen presents "
        "connection state and Seedance pack readiness in compact disclosure panels. When expanded, these "
        "panels show detected input, input/3d, output, custom-node, and workflow locations. Error messages "
        "distinguish a disconnected server, missing nodes or workflow files, a required restart, unsupported "
        "model input, generation failure, and output timeout."
    ),
    192: (
        "中文对照：网站只在技术事实会影响用户决定时显示这些事实。Device 和 PLY target 用于解释 "
        "Gaussian 运行会执行 CUDA True training 还是 initializer 转换；只有后端提供有意义的"
        "计数时，界面才显示当前训练步数。ComfyUI Video Gen 使用紧凑的折叠面板展示连接状态和 "
        "Seedance pack 准备情况。展开后，这些面板会显示检测到的 input、input/3d、output、"
        "custom-node 和 workflow 位置。错误消息会区分服务器断开、缺少 nodes 或 workflow "
        "文件、需要重启、模型输入不受支持、生成失败以及输出超时。"
    ),
    193: (
        "This approach supports trust in a long-running creative tool. A disabled training option with a reason "
        "is more useful than a button that predictably fails; a visible initializer label is more honest than "
        "presenting every PLY as trained 3DGS; and a ComfyUI readiness check is more useful than waiting for a "
        "remote-looking generation failure caused by a local folder or custom node. The website simplifies "
        "operation while retaining the minimum technical explanation needed to interpret output quality, "
        "dependency state, and recovery choices."
    ),
    194: (
        "中文对照：这一方法有助于建立用户对长时间运行创意工具的信任。与其提供一个必然失败的"
        "训练按钮，不如显示带原因的禁用选项；与其把每个 PLY 都表现为训练过的 3DGS，不如使用"
        "可见的 initializer 标签；与其等待一次看似远程、实则由本地目录或 custom node 引起的"
        "生成失败，不如预先检查 ComfyUI 准备状态。网站在简化操作的同时，仍保留解释输出质量、"
        "依赖状态和恢复选项所需的最少技术信息。"
    ),
    198: (
        "Fig. 5.1. Current sidebar showing category-sorted nodes, including ComfyUI Video Gen in the output "
        "category, together with Assets and Workflows navigation."
    ),
    199: (
        "Fig. 5.1 中文图注：当前侧栏展示按类别排序的节点，其中 ComfyUI Video Gen 位于 output "
        "类别，同时显示 Assets 与 Workflows 导航。"
    ),
    200: (
        "Fig. 5.2. Placeholder: Insert the backend architecture including frames, masks, COLMAP, Gaussian "
        "generation, mesh processing, local ComfyUI path detection, Seedance workflow submission, and video retrieval."
    ),
    201: (
        "Fig. 5.2 需要展示的图片：展示 frames、masks、COLMAP、Gaussian 生成、mesh 处理、本地 "
        "ComfyUI 路径检测、Seedance workflow 提交以及视频回收。"
    ),
    209: (
        "Control form follows parameter meaning. Frame count uses a compact numeric field, training budget uses "
        "a ten-segment bar from 1,000 to 10,000 steps, and Fast Initializer/True training uses a segmented mode "
        "selector. ComfyUI Video Gen uses text fields for addresses and prompts, selects for constrained "
        "resolution, ratio, background and render choices, numeric inputs for duration, seed, elevation and "
        "padding, and checkboxes for audio, watermark and force-render options. Connection and installation "
        "details are grouped into disclosure panels so diagnostics remain available without dominating the node."
    ),
    210: (
        "中文对照：控件形式遵循参数含义。Frame count 使用紧凑数字字段，训练预算使用从 1,000 到 "
        "10,000 steps 的十段控制条，Fast Initializer/True training 使用分段模式选择器。"
        "ComfyUI Video Gen 使用文本字段输入地址与 prompt，使用 selects 选择受约束的分辨率、"
        "比例、背景和渲染选项，使用数字输入设置时长、seed、elevation 与 padding，并使用 "
        "checkboxes 控制音频、watermark 和 force-render。连接与安装详情被组织到折叠面板中，"
        "使诊断信息保持可用，同时不会主导整个节点。"
    ),
    211: (
        "The Gaussian control stack is ordered as explanation before adjustment: Device and PLY target appear "
        "first, followed by mode availability and training steps. The ComfyUI control stack follows a similar "
        "hierarchy. Status and preset synchronization appear before the model preview, server URL and readiness "
        "checks; frequently changed generation settings remain visible; lower-frequency render settings are "
        "placed inside a collapsed panel. This ordering supports quick repeated runs while keeping the local "
        "environment inspectable when a connection or custom-node problem occurs."
    ),
    212: (
        "中文对照：Gaussian 控件栈按照“先解释、后调整”的顺序组织：Device 与 PLY target 首先"
        "出现，之后是模式可用性和训练步数。ComfyUI 控件栈采用相似层级。状态与 preset 同步位于"
        "模型预览、服务器 URL 和准备状态检查之前；经常修改的生成设置保持可见；使用频率较低的"
        "渲染设置放在折叠面板中。这样的顺序既支持快速重复运行，也能在连接或 custom-node 问题"
        "发生时检查本地环境。"
    ),
    226: (
        "Workflow feedback uses several scales. The top bar shows global Run or Stop and completed-node "
        "progress. Each node uses a compact status dot, preview-state content, and stage-specific text. "
        "Gaussian Splat Gen reports preparation, COLMAP, foreground masking, sparse or dense reconstruction, "
        "dataset preparation, Splatfacto iterations, export, completion, or a specific failure. ComfyUI Video "
        "Gen reports local connection state, Seedance readiness, prompt submission, processing, completed video, "
        "or a dependency-specific error. Green and orange node outlines were removed because they duplicated "
        "the status dot and added visual noise."
    ),
    227: (
        "中文对照：工作流反馈使用多个尺度。顶部栏显示全局 Run 或 Stop，以及已完成节点进度。"
        "每个节点使用紧凑状态点、预览状态内容和阶段专用文字。Gaussian Splat Gen 会报告 "
        "preparation、COLMAP、foreground masking、sparse 或 dense reconstruction、dataset "
        "preparation、Splatfacto iterations、export、completion 或具体失败。ComfyUI Video "
        "Gen 会报告本地连接状态、Seedance 准备情况、prompt 提交、处理中、视频完成或特定依赖"
        "错误。绿色和橙色节点外框已被移除，因为它们重复表达状态并增加视觉噪声。"
    ),
    228: (
        "Recovery actions remain close to the affected scope. A preview-level X clears one file, Clear resets "
        "all website runtime content, Stop interrupts tracked local reconstruction work, and Save Workflow "
        "stores reusable topology. ComfyUI adds Check, Install, restart-required, and Sync preset actions near "
        "the dependency they affect. Because an accepted ComfyUI prompt can outlive the website request, the "
        "interface must not imply that global Stop has definitely removed it from the external queue. This "
        "combination follows visibility-of-system-status and user-control principles [5]."
    ),
    229: (
        "中文对照：恢复操作尽量靠近受到影响的范围。预览级 X 只清空一个文件，Clear 重置网站的"
        "全部运行内容，Stop 中断被追踪的本地重建工作，Save Workflow 保存可复用拓扑。ComfyUI "
        "则把 Check、Install、restart-required 和 Sync preset 操作放在它们所影响的依赖附近。"
        "由于已被接收的 ComfyUI prompt 可能比网站请求存续更久，界面不能暗示全局 Stop 已经"
        "肯定把它从外部队列中移除。这套组合符合系统状态可见性与用户控制原则 [5]。"
    ),
    240: (
        "The data contract follows representation changes. Gaussian Splat Gen returns splatUrl or sourcePlyUrl "
        "with compute metadata. Mesh Gen converts the connected source into a conventional modelUrl and can "
        "also expose geometry_graph_surface mesh-region files produced after reconstruction. Model Cleanup "
        "and Surface Processing forward the selected model output and lightParams. ComfyUI Video Gen receives "
        "modelUrl through model-input, combines it with node settings, and returns videoUrl through video-output. "
        "Video Preview consumes that URL without receiving unrelated model or Gaussian fields."
    ),
    241: (
        "中文对照：数据契约会跟随表示变化。Gaussian Splat Gen 返回 splatUrl 或 sourcePlyUrl，"
        "并附带计算 metadata。Mesh Gen 把已连接数据源转换为传统 modelUrl，也可以公开在重建后"
        "生成的 geometry_graph_surface mesh 区域文件。Model Cleanup 与 Surface Processing "
        "会继续转发所选模型输出和 lightParams。ComfyUI Video Gen 通过 model-input 接收 "
        "modelUrl，把它与节点设置组合，并通过 video-output 返回 videoUrl。Video Preview 只"
        "消费这一 URL，不会接收无关的模型或 Gaussian 字段。"
    ),
    252: "7.4 ComfyUI API Integration",
    253: (
        "The ComfyUI integration separates preset definition, server communication, installation checks, and "
        "node presentation. A bundled API-format workflow defines Seedance3DModelLoader, multi-view rendering, "
        "the Seedance reference-generation node, preview, and SaveVideo. The preset module extracts defaults "
        "from that workflow so the website does not maintain a second conflicting parameter set. Before a run, "
        "the node can query system statistics, infer the ComfyUI base, input, input/3d and output folders, check "
        "required custom-node types and workflow files, and copy the bundled Seedance pack into detected local "
        "folders when the user selects Install."
    ),
    254: (
        "中文对照：ComfyUI 接入把 preset 定义、服务器通信、安装检查和节点呈现分开。内置的 API "
        "格式 workflow 定义了 Seedance3DModelLoader、多视角渲染、Seedance reference "
        "generation 节点、preview 和 SaveVideo。Preset 模块会从该 workflow 中提取默认值，"
        "避免网站维护第二套互相冲突的参数。运行前，节点可以查询 system statistics，推断 "
        "ComfyUI base、input、input/3d 和 output 文件夹，检查必需的 custom-node 类型与 "
        "workflow 文件，并在用户选择 Install 时把内置 Seedance pack 复制到检测出的本地文件夹。"
    ),
    255: (
        "During generation, the API validates the ephemeral session and supports GLB, GLTF, OBJ, FBX, and "
        "BLEND inputs. It copies the model into ComfyUI input/3d, injects the copied filename and current node "
        "settings into a cloned prompt, and submits the prompt with a new client identifier. The server polls "
        "ComfyUI history for up to forty-five minutes, locates an MP4, WebM, MOV, or GIF output, retrieves it "
        "through /view, and writes a session-owned copy under comfy-videos. The response returns videoUrl, "
        "videoName, promptId, and detected directory information for the node and Video Preview."
    ),
    256: (
        "中文对照：生成期间，API 会验证 ephemeral session，并支持 GLB、GLTF、OBJ、FBX 和 "
        "BLEND 输入。它把模型复制到 ComfyUI input/3d，把复制后的文件名与当前节点设置注入克隆"
        "出的 prompt，再使用新的 client identifier 提交。服务器会在最长四十五分钟内轮询 "
        "ComfyUI history，定位 MP4、WebM、MOV 或 GIF 输出，通过 /view 获取文件，并把属于"
        "当前会话的副本写入 comfy-videos。响应会返回 videoUrl、videoName、promptId 和检测到"
        "的目录信息，供节点与 Video Preview 使用。"
    ),
    258: (
        "ModelViewer, InteractiveModelViewer, and SplatViewer use event-driven render scheduling. Loader "
        "completion, controls, resize observations, material or light updates, and mesh-region highlighting "
        "request a frame; cleanup removes listeners and disposes geometries, materials, renderers, and animation "
        "handles. ComfyUI Video Gen does not create another live Three.js scene. It presents lightweight status "
        "and model-readiness information, while final playback is delegated to Video Preview. This keeps the "
        "larger output node from adding persistent GPU work to the editor."
    ),
    259: (
        "中文对照：ModelViewer、InteractiveModelViewer 和 SplatViewer 使用事件驱动的渲染调度。"
        "加载完成、controls、resize observations、材质或灯光更新以及 mesh 区域高亮会请求新帧；"
        "清理阶段则移除 listeners，并释放 geometries、materials、renderers 和 animation "
        "handles。ComfyUI Video Gen 不会创建另一套实时 Three.js 场景，而是显示轻量状态和模型"
        "准备信息；最终播放交给 Video Preview。这样可以避免较大的输出节点给编辑器增加持续 GPU 工作。"
    ),
    260: (
        "Verification combines static, scripted, and browser checks. TypeScript checking validates node "
        "contracts and API result shapes. Existing reconstruction tests cover foreground masks, point filtering, "
        "Gaussian routing, task polling, model cleanup, and Clear behaviour. ComfyUI tests cover preset extraction, "
        "prompt replacement, output discovery, folder detection, Seedance pack readiness and installation paths, "
        "default-workflow edges, model-to-video data transfer, and terminal-node completion. Browser inspection "
        "confirms current Gaussian controls, ComfyUI disclosure panels, Surface Processing handoff, Video Preview "
        "delivery, and low idle rendering activity."
    ),
    261: (
        "中文对照：验证工作结合了静态检查、脚本检查和浏览器检查。TypeScript 检查用于验证节点契约"
        "和 API 结果结构。现有重建测试覆盖前景遮罩、点过滤、Gaussian routing、task polling、"
        "model cleanup 和 Clear 行为。ComfyUI 测试覆盖 preset 提取、prompt 替换、输出发现、"
        "文件夹检测、Seedance pack 准备状态与安装路径、默认工作流 edges、模型到视频数据传输和"
        "末端节点完成判断。浏览器检查会确认当前 Gaussian 控件、ComfyUI 折叠面板、Surface "
        "Processing 交接、Video Preview 传递以及较低的空闲渲染活动。"
    ),
    262: (
        "Several safeguards remain contract-based rather than visual. Original frames, masks, COLMAP "
        "registrations, and Nerfstudio entries must preserve matching stems. Cancellation checks around expensive "
        "reconstruction stages prevent a stopped run from silently starting the next local script. ComfyUI uses "
        "a separate process and queue, so website cancellation currently stops waiting and downstream state but "
        "does not call a dedicated ComfyUI interrupt endpoint. This distinction is included in testing notes and "
        "future work rather than being concealed by the global Stop label."
    ),
    263: (
        "中文对照：若干保护仍依赖数据契约，而不是可见界面。原始帧、masks、COLMAP registrations "
        "和 Nerfstudio entries 必须保持匹配的文件名主体。在昂贵重建阶段前后执行取消检查，可以"
        "防止已停止的运行悄悄启动下一个本地脚本。ComfyUI 使用独立进程和队列，因此网站取消目前"
        "只会停止等待和下游状态，而不会调用专用 ComfyUI interrupt endpoint。报告把这一区分"
        "写入测试说明与未来工作，而不会用全局 Stop 标签将其隐藏。"
    ),
    264: (
        "Fig. 7.1. Current Gaussian Splat Gen showing detected Device, PLY target, Fast Initializer/True "
        "training availability, and the segmented training-step control."
    ),
    265: (
        "Fig. 7.1 中文图注：当前 Gaussian Splat Gen，展示检测到的 Device、PLY target、Fast "
        "Initializer/True training 可用性以及训练步数分段控制条。"
    ),
    270: (
        "The prototype now presents a coherent path from video capture or direct PLY input to reusable splat, "
        "mesh, surface-processed model, and ComfyUI-generated video outputs. Its strongest result is explainability: "
        "connected ports reveal data flow; Device and PLY target reveal the compute route; real progress reflects "
        "backend work; Stop reaches tracked local process trees; Clear removes stale previews; and the protected "
        "Default Workflow restores a known baseline. The new output node makes local ComfyUI status, Seedance "
        "readiness, generation settings, prompt identifier, and final video delivery visible in the same graph."
    ),
    271: (
        "中文对照：原型现在提供了一条从视频拍摄或直接 PLY 输入，到可复用 splat、mesh、表面"
        "处理模型和 ComfyUI 生成视频输出的一致路径。其最突出的成果是可解释性：已连接端口展示"
        "数据流；Device 与 PLY target 展示计算路径；真实进度反映后端工作；Stop 能到达被追踪"
        "的本地进程树；Clear 会移除残留预览；受保护的 Default Workflow 可以恢复已知基线。"
        "新的输出节点还在同一张节点图中公开本地 ComfyUI 状态、Seedance 准备情况、生成设置、"
        "prompt identifier 和最终视频传递。"
    ),
    274: "8.2 Reconstruction Quality and Video-Generation Limits",
    277: (
        "ComfyUI video output introduces a different form of uncertainty from reconstruction. A technically "
        "valid GLB or FBX can be submitted successfully while the generated clip still changes small jewellery "
        "details, introduces inconsistent reflections, selects an unsuitable camera path, or produces a result "
        "that depends strongly on prompt wording and reference rendering. The current integration verifies "
        "transport, preset substitution, dependency readiness, and output recovery; it does not prove visual "
        "fidelity to the source model. Final evaluation should therefore compare shape, distinctive details, "
        "materials, camera motion, and temporal stability rather than treating file generation alone as success."
    ),
    278: (
        "中文对照：ComfyUI 视频输出带来的不确定性不同于重建。技术上有效的 GLB 或 FBX 可以成功"
        "提交，但生成片段仍可能改变细小首饰细节、产生不一致反射、选择不合适的相机路径，或得到"
        "高度依赖 prompt 写法与参考渲染的结果。当前接入验证的是传输、preset 替换、依赖准备状态"
        "和输出回收，并不能证明视频在视觉上忠实于源模型。因此，最终评估应比较形状、显著细节、"
        "材质、相机运动和时间稳定性，而不能只把文件成功生成视为成功。"
    ),
    279: (
        "Product category changes the likely failure mode. Rings and jewellery require preservation of stone "
        "shape, prong count, band profile, and reflective metal; cosmetics packaging requires label, cap, colour, "
        "and proportion consistency; furniture requires stable large-scale geometry and plausible camera clearance. "
        "The bundled Seedance preset is deliberately product-oriented, but one prompt and one multi-view render "
        "strategy cannot guarantee equal performance across categories. Category-specific presets and evaluation "
        "criteria are therefore more defensible than a universal quality claim."
    ),
    280: (
        "中文对照：产品类别会改变最可能出现的失败模式。戒指和首饰要求保留宝石形状、镶爪数量、"
        "戒圈轮廓和反光金属；化妆品包装要求标签、瓶盖、颜色和比例保持一致；家具则要求大尺度"
        "几何稳定，并具有合理的相机避让。内置 Seedance preset 有意面向产品，但同一个 prompt "
        "和同一种多视角渲染策略不能保证在所有类别上表现一致。因此，针对类别的 preset 和评估"
        "标准比通用质量声明更可靠。"
    ),
    282: (
        "The project deliberately favours an integrated workflow over exposing every COLMAP, Open3D, Blender, "
        "Nerfstudio, or ComfyUI parameter. This reduces command-line and file-management burden but limits expert "
        "tuning. Fast Initializer allows Apple Silicon and CPU users to complete the workflow, yet its PLY is not "
        "equivalent to CUDA Splatfacto training. Mesh conversion enables GLB editing and ComfyUI model input but "
        "loses view-dependent splat appearance. ComfyUI extends the prototype from reconstruction to presentation "
        "video, while introducing a second local service, custom-node dependencies, possible model/API cost, and "
        "a queue that is not yet controlled by the website Stop action."
    ),
    283: (
        "中文对照：项目有意优先选择整合工作流，而不是公开 COLMAP、Open3D、Blender、Nerfstudio "
        "或 ComfyUI 的每一个参数。这样可以降低命令行和文件管理负担，但也限制专家级调节。Fast "
        "Initializer 允许 Apple Silicon 与 CPU 用户完成工作流，但其 PLY 并不等同于 CUDA "
        "Splatfacto 训练。Mesh 转换支持 GLB 编辑和 ComfyUI 模型输入，却会丢失依赖视角的 splat "
        "外观。ComfyUI 把原型从重建扩展到展示视频，同时引入第二项本地服务、custom-node 依赖、"
        "可能的模型/API 成本，以及尚未受到网站 Stop 操作控制的队列。"
    ),
    284: (
        "Temporarily disabling Material Gen remains an explicit trade-off. It reduces external cost, credential "
        "risk, and default-workflow failure but removes text-to-texture generation from the current prototype. "
        "ComfyUI Video Gen is retained in the Default Workflow because its dependency can be inspected through "
        "connection, folder, pack, and restart checks. This does not make the dependency invisible or guaranteed; "
        "it makes it diagnosable. The contrast illustrates the project's design principle: an external capability "
        "belongs in the default experience only when its state and failure boundary can be communicated."
    ),
    285: (
        "中文对照：暂时停用 Material Gen 仍是一项明确取舍。它降低了外部成本、凭证风险和默认"
        "工作流失败概率，但也从当前原型中移除了 text-to-texture 生成。ComfyUI Video Gen "
        "保留在 Default Workflow 中，是因为可以通过连接、文件夹、pack 和重启检查来观察其"
        "依赖状态。这并不意味着依赖被隐藏或得到保证，而是让它可以被诊断。二者的对比说明项目"
        "的一项设计原则：只有当外部能力的状态和失败边界可以被传达时，它才适合进入默认体验。"
    ),
    287: (
        "Against the capstone objectives, the system succeeds as a design-led integration prototype. It defines "
        "typed representation boundaries, runs a connected graph, supports two Gaussian input routes, exposes "
        "hardware capability, manages cancellable local reconstruction tasks, reuses assets and workflows, "
        "preserves model and lighting context, connects the final model to local ComfyUI, and forwards the "
        "generated video to a dedicated preview. The Figma kit records the corresponding visual system and the "
        "website keeps dependencies visible through controls, status, and errors."
    ),
    288: (
        "中文对照：对照毕业设计目标，系统已经作为以设计为主导的整合原型取得成果。它定义了"
        "带类型的表示边界，运行已连接节点图，支持两条 Gaussian 输入路径，公开硬件能力，管理"
        "可取消的本地重建任务，复用资产和工作流，保留模型与灯光上下文，把最终模型连接到本地 "
        "ComfyUI，并把生成视频转发到专用预览。Figma kit 记录相应视觉系统，网站则通过控件、"
        "状态和错误让依赖保持可见。"
    ),
    289: (
        "The system remains incomplete as a production 3DGS and media platform. True training still depends on "
        "CUDA-compatible Nerfstudio and gsplat, foreground-mask quality has not been benchmarked, large splat "
        "rendering is approximate, mesh reconstruction can produce defects, and external tools must be installed "
        "on the server. ComfyUI generation depends on a separately running environment and has no website-driven "
        "queue interrupt in the current version. These limitations define the boundary between a functioning "
        "capstone workflow and a production service. Every result should be labelled with its input route, "
        "representation, hardware, and generation backend."
    ),
    290: (
        "中文对照：系统仍不是完整的生产级 3DGS 与媒体平台。True training 依然依赖兼容 CUDA 的 "
        "Nerfstudio 和 gsplat；前景遮罩质量尚未经过基准测试；大型 splat 渲染只是近似效果；"
        "mesh 重建可能产生缺陷；服务器还必须安装外部工具。ComfyUI 生成依赖单独运行的环境，"
        "当前版本也没有由网站驱动的队列中断。这些限制明确了可运行毕业设计工作流与生产服务之间"
        "的边界。每项结果都应标明其输入路径、表示类型、硬件和生成后端。"
    ),
    297: (
        "This capstone project designed and implemented Splat to Sculpt as a capability-aware node workflow for "
        "Gaussian reconstruction, mesh production, surface processing, and local ComfyUI video generation. The "
        "workflow distinguishes frames, point clouds, initializer splats, trained splats, meshes, GLB assets, "
        "surface-processed models, and generated videos through explicit ports and conversion nodes. It supports "
        "video and direct-PLY input, CUDA True training when available, MPS/CPU Fast Initializer, foreground masks, "
        "multi-view point filtering, cancellable local tasks, asset reuse, model previews, ComfyUI dependency checks, "
        "Seedance preset synchronization, and final video preview."
    ),
    298: (
        "中文对照：本毕业设计把 Splat to Sculpt 设计并实现为具有设备能力判断的节点式工作流，"
        "用于 Gaussian 重建、mesh 制作、表面处理和本地 ComfyUI 视频生成。工作流通过显式端口"
        "与转换节点区分 frames、point clouds、initializer splats、trained splats、meshes、"
        "GLB assets、表面处理模型和生成视频。系统支持视频与直接 PLY 输入、可用时的 CUDA True "
        "training、MPS/CPU Fast Initializer、前景遮罩、多视角点过滤、可取消本地任务、资产复用、"
        "模型预览、ComfyUI 依赖检查、Seedance preset 同步和最终视频预览。"
    ),
    302: (
        "The project shows that interface claims must be backed by system behaviour. Stop is credible only for "
        "processes the application can actually terminate; progress is credible only when it follows real stages; "
        "True training is credible only when images, cameras, and a compatible runtime exist; and a generated video "
        "is credible only when its local dependency, prompt submission, output recovery, and limitations are visible. "
        "The same principle applies to performance: a completed workflow should not leave hidden render loops "
        "consuming the GPU."
    ),
    303: (
        "中文对照：项目表明，界面中的声明必须由系统行为支持。只有对于应用能够真正终止的进程，"
        "Stop 才可信；只有当进度跟随真实阶段时，进度才可信；只有在图像、相机和兼容运行环境都"
        "存在时，True training 才可信；只有当本地依赖、prompt 提交、输出回收和限制都保持可见"
        "时，生成视频才可信。相同原则也适用于性能：完成的工作流不应留下隐藏渲染循环继续消耗 GPU。"
    ),
    304: (
        "The most valuable outcome is therefore not one reconstruction or video model but a product layer that "
        "exposes choices and limitations without forcing users to operate every tool directly. Foreground filtering, "
        "initializer routing, on-demand preview, typed model-to-video handoff, ComfyUI readiness checks, and the "
        "design system all emerged from observed failures or ambiguity. They demonstrate how debugging evidence "
        "can become workflow and UI design knowledge."
    ),
    305: (
        "中文对照：因此，最有价值的成果不是某一种重建或视频模型，而是一个产品层：它公开选择"
        "和限制，同时不强迫用户直接操作所有工具。Foreground filtering、initializer routing、"
        "按需预览、带类型的模型到视频交接、ComfyUI 准备状态检查和设计系统，都源于实际观察到的"
        "失败或歧义。它们说明调试证据如何转化为工作流与 UI 设计知识。"
    ),
    308: (
        "Future work should first evaluate reconstruction and video generation quantitatively across jewellery, "
        "cosmetics packaging, furniture, reflective objects, and cluttered backgrounds. Reconstruction metrics "
        "could compare registration rate, retained background points, completeness, and time. Video evaluation "
        "should compare product-shape fidelity, distinctive-part preservation, material consistency, temporal "
        "stability, and prompt sensitivity. The ComfyUI route should add prompt cancellation through an interrupt "
        "or queue API, detect completion through WebSocket events where appropriate, clean copied input files, "
        "and report queue position and remaining dependency costs."
    ),
    309: (
        "中文对照：未来工作首先应针对首饰、化妆品包装、家具、反光物体和杂乱背景，定量评估"
        "重建与视频生成。重建指标可以比较注册率、残留背景点、完整度和时间；视频评估应比较产品"
        "形状忠实度、显著部件保留、材质一致性、时间稳定性和 prompt 敏感度。ComfyUI 路由应通过 "
        "interrupt 或 queue API 加入 prompt 取消，在适当情况下使用 WebSocket 事件检测完成，"
        "清理复制的输入文件，并报告队列位置和剩余依赖成本。"
    ),
    310: (
        "A complete browser splat renderer remains another priority, including anisotropic covariance projection, "
        "correct alpha compositing or tile sorting, view-dependent spherical harmonics, progressive loading, and "
        "level of detail. Remote CUDA execution could allow Mac users to keep the same node interface while sending "
        "Splatfacto training to a managed GPU queue. A deployable ComfyUI service would likewise need authentication, "
        "storage quotas, model licensing, API-cost policy, concurrency limits, monitoring, and isolation from user "
        "files. Material Gen should be restored only after selecting a stable service and defining comparable policies."
    ),
    311: (
        "中文对照：完整的浏览器 splat renderer 仍是另一项优先工作，包括各向异性协方差投影、"
        "正确的 alpha compositing 或 tile sorting、依赖视角的 spherical harmonics、"
        "progressive loading 和 level of detail。远程 CUDA 执行可以让 Mac 用户继续使用相同"
        "节点界面，同时把 Splatfacto 训练发送到受管理的 GPU 队列。可部署的 ComfyUI 服务同样"
        "需要 authentication、storage quotas、模型许可、API 成本策略、并发限制、监测，以及与"
        "用户文件的隔离。Material Gen 只有在选定稳定服务并定义同类策略后才应恢复。"
    ),
    313: (
        "Product development should add capture guidance and reconstruction diagnostics before training begins. "
        "Registered-image count, camera coverage, foreground-mask confidence, sparse-point count, and warnings for "
        "blur, reflection, or insufficient overlap would help users decide whether to continue. Before ComfyUI "
        "submission, a preflight summary should verify model format, file size, detected input folder, required "
        "nodes, workflow availability, selected model, estimated duration or cost, and cancellation boundary. "
        "Larger workflows would also benefit from node groups, subgraphs, automatic layout, and reusable templates."
    ),
    314: (
        "中文对照：产品开发应在训练开始前加入拍摄指导和重建诊断。已注册图像数量、相机覆盖、"
        "前景遮罩置信度、稀疏点数量，以及对模糊、反光或重叠不足的警告，都可以帮助用户决定是否"
        "继续。在提交 ComfyUI 之前，preflight summary 应验证模型格式、文件大小、检测到的 input "
        "文件夹、必需 nodes、workflow 可用性、所选模型、预计时长或成本，以及取消边界。较大"
        "工作流还可以从 node groups、subgraphs、automatic layout 和可复用 templates 中受益。"
    ),
    326: (
        "• Figures 1.1, 4.1, 5.1, 6.1–6.3, 7.1, and 7.3 use current website or Figma images. "
        "Replace the remaining concept and experimental placeholders with final diagrams and evaluated outputs."
    ),
    327: (
        "中文对照：图 1.1、4.1、5.1、6.1–6.3、7.1 和 7.3 已使用当前网页或 Figma 图片。"
        "其余概念图与实验占位区应在最终提交前替换为完成的图示和经过评估的输出。"
    ),
    332: (
        "• Label CUDA Splatfacto, MPS/CPU initializer, direct-PLY initializer, mesh reconstruction, "
        "and ComfyUI/Seedance video results accurately in every caption."
    ),
    333: (
        "中文对照：请在每条图注中准确标明 CUDA Splatfacto、MPS/CPU initializer、直接 PLY "
        "initializer、mesh 重建以及 ComfyUI/Seedance 视频结果。"
    ),
}


FIGURE_SPECS = {
    "1.1": {
        "table": 1,
        "image": "fig_1_1_editor.png",
        "width": 6.15,
        "alt": "Current Splat to Sculpt editor with ComfyUI Video Gen in the connected workflow.",
        "list": (
            "Fig. 1.1: Current Splat to Sculpt editor and connected ComfyUI workflow."
        ),
    },
    "4.1": {
        "table": 6,
        "image": "fig_4_1_workflow.png",
        "width": 6.2,
        "alt": "Current Default Workflow ending with ComfyUI Video Gen and Video Preview.",
        "list": (
            "Fig. 4.1: Current Default Workflow from video input through Gaussian and mesh processing "
            "to ComfyUI Video Gen and Video Preview."
        ),
    },
    "5.1": {
        "table": 8,
        "image": "fig_5_1_sidebar_assets.png",
        "width": 3.35,
        "alt": "Current node library including ComfyUI Video Gen in the output category.",
        "list": (
            "Fig. 5.1: Current category-sorted node library including ComfyUI Video Gen."
        ),
    },
    "7.1": {
        "table": 13,
        "image": "fig_7_1_gaussian_controls.png",
        "width": 3.65,
        "alt": "Current Gaussian Splat Gen controls without the removed point-cloud layering option.",
        "list": (
            "Fig. 7.1: Current Gaussian Splat Gen Device, PLY target, training mode, and training-step controls."
        ),
    },
}


LIST_REPLACEMENTS = {
    "Fig. 2.1:": (
        "Fig. 2.1: End-to-end data flow from video or PLY input through Gaussian generation, "
        "mesh conversion, Surface Processing, ComfyUI Video Gen, and Video Preview."
    ),
    "Fig. 3.1:": (
        "Fig. 3.1: User journey from capture or asset reuse through reconstruction, refinement, "
        "ComfyUI generation, and final video inspection."
    ),
    "Fig. 4.2:": (
        "Fig. 4.2: Frame/COLMAP and direct-PLY Gaussian routes with the shared downstream "
        "model-to-ComfyUI video handoff."
    ),
    "Fig. 5.2:": (
        "Fig. 5.2: Backend architecture including reconstruction tools, local ComfyUI path detection, "
        "Seedance workflow submission, and video retrieval."
    ),
    "Fig. 7.1:": FIGURE_SPECS["7.1"]["list"],
    "Fig. 8.1:": (
        "Fig. 8.1: Before-and-after comparison of the earlier workflow and the current "
        "typed workflow with local ComfyUI output."
    ),
    "Fig. 8.2:": (
        "Fig. 8.2: Representative masks, filtered point cloud, splat, mesh, and "
        "ComfyUI/Seedance video result."
    ),
}


def clone_run_properties(source_run, target_run) -> None:
    if source_run is None or source_run._r.rPr is None:
        return
    target_run._r.insert(0, deepcopy(source_run._r.rPr))


def replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
    source_run = paragraph.runs[0] if paragraph.runs else None
    paragraph.clear()
    run = paragraph.add_run(text)
    clone_run_properties(source_run, run)


def insert_paragraph_after(source: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    if source._p.pPr is not None:
        new_p.append(deepcopy(source._p.pPr))
    source._p.addnext(new_p)
    paragraph = Paragraph(new_p, source._parent)
    run = paragraph.add_run(text)
    if source.runs:
        clone_run_properties(source.runs[0], run)
    return paragraph


def insert_paragraph_before(source: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    source._p.addprevious(new_p)
    paragraph = Paragraph(new_p, source._parent)
    if text:
        paragraph.add_run(text)
    return paragraph


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def set_cell_margins(cell, margin: int = 45) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge in ("top", "start", "bottom", "end"):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(margin))
        element.set(qn("w:type"), "dxa")


def add_alt_text(run, text: str) -> None:
    doc_properties = run._element.xpath(".//wp:docPr")
    if doc_properties:
        doc_properties[0].set("descr", text)
        doc_properties[0].set("name", text[:80])


def replace_figure_table(document: Document, spec: dict) -> None:
    table = document.tables[spec["table"]]
    cell = table.cell(0, 0)
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, "FFFFFF")
    set_cell_margins(cell)
    table.rows[0].height = None
    table.rows[0].height_rule = None
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(FIGURES_DIR / spec["image"]), width=Inches(spec["width"]))
    add_alt_text(run, spec["alt"])


def format_caption(paragraph: Paragraph, *, chinese: bool) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = not chinese
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(8)
        run.font.italic = not chinese
        run._element.get_or_add_rPr().get_or_add_rFonts().set(
            qn("w:eastAsia"), "SimSun" if chinese else "Times New Roman"
        )


def set_table_cell(cell, text: str) -> None:
    source = cell.paragraphs[0]
    replace_paragraph_text(source, text)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)


def update_requirement_table(document: Document) -> None:
    table = document.tables[3]
    rows = [
        ["Requirement Type", "Requirement", "Design Response"],
        [
            "Functional",
            "Video-to-3D workflow",
            "Connected chain from video or PLY to splat, mesh, surface model, ComfyUI video, and preview.",
        ],
        [
            "Functional",
            "Dual Gaussian input",
            "Frames support COLMAP/True training when available; direct PLY is routed to Fast Initializer.",
        ],
        [
            "Functional",
            "Local ComfyUI output",
            "Validate a final model, check local Seedance readiness, submit the API workflow, and return videoUrl.",
        ],
        [
            "Functional",
            "Preview and reuse",
            "Visual previews, asset thumbnails, saved workflows, a protected preset, and generated-video assets.",
        ],
        [
            "Non-functional",
            "Stop and recovery",
            "Tracked local process termination, global Clear, and explicit disclosure of the external ComfyUI queue limit.",
        ],
        [
            "Non-functional",
            "Dependency transparency",
            "Visible CUDA/MPS/CPU route, ComfyUI connection, detected folders, Seedance status, and actionable errors.",
        ],
        [
            "Non-functional",
            "Idle performance",
            "On-demand WebGL rendering and complete renderer/resource cleanup.",
        ],
    ]
    for row, values in zip(table.rows, rows):
        for cell, value in zip(row.cells, values):
            set_table_cell(cell, value)


def update_node_table(document: Document) -> None:
    table = document.tables[5]
    table.add_row()
    rows = [
        ["Node", "Input", "Output", "Current Role"],
        ["Video Upload", "Video", "Server video", "Source upload and target frame count"],
        ["Frame Extraction", "Video", "Images", "Controlled temporal sampling"],
        [
            "Gaussian Splat Gen",
            "Images or PLY",
            "Splat PLY / source PLY",
            "Capability-aware True training or Fast Initializer",
        ],
        [
            "Mesh Gen",
            "Splat/PLY/OBJ/GLB",
            "GLB/OBJ/PLY",
            "Representation conversion and post-reconstruction mesh regions",
        ],
        ["Model Cleanup", "Model", "Cleaned model", "Blender cleanup and model forwarding"],
        [
            "Surface Processing",
            "Model",
            "Processed model + light params",
            "Material, colour, lighting, and Blender render application",
        ],
        [
            "ComfyUI Video Gen",
            "GLB/GLTF/OBJ/FBX/BLEND",
            "Video URL",
            "Local ComfyUI/Seedance model-to-video generation",
        ],
        ["Video Preview", "Video URL", "Playback/download", "Final generated-video inspection"],
        ["Material Gen", "Text", "PNG", "Temporarily disabled; retained for compatibility"],
        ["Sticky Note", "Text", "Annotation", "Non-computational workflow explanation"],
    ]
    for row, values in zip(table.rows, rows):
        for cell, value in zip(row.cells, values):
            set_table_cell(cell, value)


def add_comfy_figure(document: Document) -> None:
    heading = [
        p for p in document.paragraphs if p.text.strip() == "8. Results and Discussion"
    ][-1]
    image_paragraph = insert_paragraph_before(heading)
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.space_before = Pt(3)
    image_paragraph.paragraph_format.space_after = Pt(2)
    image_paragraph.paragraph_format.keep_with_next = True
    run = image_paragraph.add_run()
    run.add_picture(str(FIGURES_DIR / "fig_7_3_comfy_controls.png"), width=Inches(3.75))
    add_alt_text(
        run,
        "ComfyUI Video Gen node showing connection, Seedance readiness, and generation settings.",
    )

    en = insert_paragraph_before(
        heading,
        "Fig. 7.3. ComfyUI Video Gen status and controls.",
    )
    format_caption(en, chinese=False)
    zh = insert_paragraph_before(
        heading,
        "Fig. 7.3 中文图注：ComfyUI Video Gen 的状态与控件。",
    )
    format_caption(zh, chinese=True)

    list_72 = next(p for p in document.paragraphs if p.text.strip().startswith("Fig. 7.2:"))
    inserted = insert_paragraph_after(
        list_72,
        "Fig. 7.3: ComfyUI Video Gen controls and status.",
    )
    if list_72.runs and inserted.runs:
        clone_run_properties(list_72.runs[0], inserted.runs[0])


def remove_redundant_paragraph_pairs(document: Document) -> None:
    english_prefixes = (
        "The same transparency applies to saved and temporary resources.",
        "The manual migration process also revealed a design-system limitation:",
        "Product category changes the likely failure mode.",
        "A complete browser splat renderer remains another priority,",
    )
    for prefix in english_prefixes:
        paragraphs = document.paragraphs
        index = next(i for i, paragraph in enumerate(paragraphs) if paragraph.text.startswith(prefix))
        pair = paragraphs[index:index + 2]
        if len(pair) != 2 or not pair[1].text.startswith("中文对照："):
            raise RuntimeError(f"Missing Chinese pair for paragraph: {prefix}")
        for paragraph in reversed(pair):
            parent = paragraph._p.getparent()
            parent.remove(paragraph._p)


def normalize_cjk_fonts(document: Document) -> None:
    cjk_pattern = re.compile(r"[\u3400-\u9fff]")
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)

    for paragraph in paragraphs:
        if not cjk_pattern.search(paragraph.text):
            continue
        for run in paragraph.runs:
            fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
            fonts.set(qn("w:eastAsia"), "Arial Unicode MS")
            fonts.set(qn("w:hint"), "eastAsia")


def tidy_front_matter_pagination(document: Document) -> None:
    toc = next(
        i
        for i, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip() == "Table of Contents"
    )
    acknowledgement = next(
        paragraph
        for paragraph in document.paragraphs[:toc]
        if paragraph.text.strip() == "Acknowledgement"
    )
    acknowledgement.paragraph_format.page_break_before = True

    statement = next(
        i
        for i, paragraph in enumerate(document.paragraphs[:toc])
        if paragraph.text.strip() == "Statement of Contributions"
    )
    for paragraph in document.paragraphs[statement + 1:toc]:
        if paragraph.text.strip():
            paragraph.paragraph_format.line_spacing = 1.15
        if paragraph.text.startswith("中文对照："):
            paragraph.paragraph_format.space_after = Pt(2)


def approximate_english_word_count(document: Document) -> int:
    texts: list[str] = []
    start = next(i for i, p in enumerate(document.paragraphs) if p.text.strip() == "Abstract")
    end = [
        i for i, p in enumerate(document.paragraphs) if p.text.strip() == "References"
    ][-1]
    for paragraph in document.paragraphs[start:end]:
        text = paragraph.text.strip()
        if not text:
            continue
        if text.startswith("中文对照：") or "中文图注：" in text or "需要展示的图片：" in text:
            continue
        texts.append(text)
    for table_index in (3, 5):
        for row in document.tables[table_index].rows:
            for cell in row.cells:
                texts.append(cell.text)
    return len(re.findall(r"[A-Za-z0-9]+(?:['’-][A-Za-z0-9]+)*", " ".join(texts)))


def main() -> None:
    input_path = Path(sys.argv[1]) if len(sys.argv) > 1 else INPUT_PATH
    output_path = Path(sys.argv[2]) if len(sys.argv) > 2 else OUTPUT_PATH
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document(input_path)

    for index, text in REPLACEMENTS.items():
        replace_paragraph_text(document.paragraphs[index], text)

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        for prefix, replacement in LIST_REPLACEMENTS.items():
            if text.startswith(prefix):
                replace_paragraph_text(paragraph, replacement)
                break

    for figure_number, spec in FIGURE_SPECS.items():
        replace_figure_table(document, spec)
        for paragraph in document.paragraphs:
            if paragraph.text.strip().startswith(f"Fig. {figure_number}."):
                format_caption(paragraph, chinese=False)
            elif paragraph.text.strip().startswith(f"Fig. {figure_number} 中文图注"):
                format_caption(paragraph, chinese=True)

    for figure_number, spec in FIGURE_SPECS.items():
        for paragraph in document.paragraphs:
            if paragraph.text.strip().startswith(f"Fig. {figure_number}:"):
                replace_paragraph_text(paragraph, spec["list"])

    update_requirement_table(document)
    update_node_table(document)
    add_comfy_figure(document)
    remove_redundant_paragraph_pairs(document)
    normalize_cjk_fonts(document)
    tidy_front_matter_pagination(document)

    toc_heading_index = next(
        i for i, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip() == "Table of Contents"
    )
    toc_after = [
        paragraph.text.strip()
        for paragraph in document.paragraphs[toc_heading_index + 1:toc_heading_index + 18]
    ]
    if toc_after != EXPECTED_TOC:
        raise RuntimeError(
            "Table of Contents text changed unexpectedly: "
            f"expected={EXPECTED_TOC!r}, actual={toc_after!r}"
        )

    count = approximate_english_word_count(document)
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith("Approximate English word count in report body and captions:"):
            replace_paragraph_text(
                paragraph,
                f"Approximate English word count in report body and captions: {count:,} words.",
            )
        elif paragraph.text.strip().startswith("中文对照：报告正文和图注中的英文词数约为"):
            replace_paragraph_text(
                paragraph,
                f"中文对照：报告正文和图注中的英文词数约为 {count:,} 词。",
            )

    document.save(output_path)
    print(output_path)
    print(f"approximate_english_word_count={count}")


if __name__ == "__main__":
    main()
