from pathlib import Path
import sys

from docx import Document


def main() -> None:
    path = Path(sys.argv[1])
    document = Document(path)
    image_relations = sum(
        1 for relation in document.part.rels.values() if "image" in relation.reltype
    )
    print(
        f"paragraphs={len(document.paragraphs)} tables={len(document.tables)} "
        f"inline_shapes={len(document.inline_shapes)} image_relations={image_relations}"
    )
    for table_index, table in enumerate(document.tables):
        table_text = " | ".join(
            " ".join(cell.text.split())
            for row in table.rows
            for cell in row.cells
        )
        print(
            f"table {table_index}: rows={len(table.rows)} cols={len(table.columns)} "
            f"text={table_text[:300]!r}"
        )
    keywords = ("fig.", "figure", "placeholder", "截图", "占位")
    for index, paragraph in enumerate(document.paragraphs):
        text = " ".join(paragraph.text.split())
        if any(keyword in text.lower() for keyword in keywords):
            print(f"{index}: {text}")

    if len(sys.argv) > 2:
        start, end = (int(value) for value in sys.argv[2].split(":", 1))
        print(f"-- paragraphs {start}:{end} --")
        for index in range(start, min(end, len(document.paragraphs))):
            paragraph = document.paragraphs[index]
            text = " ".join(paragraph.text.split())
            print(
                f"{index}: style={paragraph.style.name!r} "
                f"alignment={paragraph.alignment!r} text={text!r} "
                f"drawings={len(paragraph._p.xpath('.//w:drawing'))} "
                f"picts={len(paragraph._p.xpath('.//w:pict'))} "
                f"all_text={' | '.join(paragraph._p.xpath('.//w:t/text()'))!r}"
            )


if __name__ == "__main__":
    main()
