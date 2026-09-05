from copy import deepcopy
from pathlib import Path
import sys

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


FIGURES = {
    "1.1": {
        "table": 1,
        "image": "fig_1_1_editor.png",
        "width": 6.15,
        "en": (
            "Fig. 1.1. Current Splat to Sculpt editor showing the top bar, "
            "category-sorted node library, Gaussian controls, and connected canvas."
        ),
        "zh": (
            "Fig. 1.1 当前的 Splat to Sculpt 编辑器，展示顶部控制栏、按类别排序的节点库、"
            "Gaussian 控制区以及已连线的画布。"
        ),
        "list": (
            "Fig. 1.1: Current Splat to Sculpt editor with the top bar, node library, "
            "Gaussian controls, and connected canvas."
        ),
    },
    "4.1": {
        "table": 6,
        "image": "fig_4_1_workflow.png",
        "width": 6.2,
        "en": (
            "Fig. 4.1. Current workflow graph in the editor, showing the connected "
            "sequence from Gaussian Splat Gen through Mesh Gen, cleanup, surface "
            "processing, and final preview."
        ),
        "zh": (
            "Fig. 4.1 编辑器中的当前工作流图，展示从 Gaussian Splat Gen 经过 Mesh Gen、"
            "模型清理和表面处理，最终到达预览输出的连线顺序。"
        ),
        "list": (
            "Fig. 4.1: Current connected workflow from Gaussian Splat Gen through "
            "mesh processing and final preview."
        ),
    },
    "5.1": {
        "table": 8,
        "image": "fig_5_1_sidebar_assets.png",
        "width": 5.0,
        "en": (
            "Fig. 5.1. Sidebar navigation and asset-library views. The Nodes view "
            "groups available tools, while the Assets view presents videos, splats, "
            "models, and rendered videos with identifying thumbnails."
        ),
        "zh": (
            "Fig. 5.1 侧栏导航与资产库视图。Nodes 视图按类别组织可用工具，Assets 视图则通过"
            "识别缩略图展示视频、splat、模型和渲染视频。"
        ),
        "list": (
            "Fig. 5.1: Sidebar Nodes and Assets views with category grouping and "
            "identifying asset thumbnails."
        ),
    },
    "6.1": {
        "table": 10,
        "image": "fig_6_1_node_cards.png",
        "width": 6.2,
        "en": (
            "Fig. 6.1. Figma node-card component overview showing consistent headers, "
            "handles, status dots, preview boxes, controls, and delete actions across "
            "the available node types."
        ),
        "zh": (
            "Fig. 6.1 Figma 节点卡片组件总览，展示各类节点一致的标题栏、端口、状态点、"
            "预览框、控制项和删除操作。"
        ),
        "list": (
            "Fig. 6.1: Figma node-card component overview across the available node types."
        ),
    },
    "6.2": {
        "table": 11,
        "image": "fig_6_2_ui_kit.png",
        "width": 6.2,
        "en": (
            "Fig. 6.2. Figma UI Kit component variants for top-bar actions, buttons, "
            "training steps, preview states, sidebar/asset/workflow rows, and node cards."
        ),
        "zh": (
            "Fig. 6.2 Figma UI Kit 的组件 variants，涵盖顶部栏操作、按钮、训练步数、"
            "预览状态、侧栏/资产/工作流条目以及节点卡片。"
        ),
        "list": (
            "Fig. 6.2: Figma UI Kit variants for controls, previews, sidebar rows, "
            "assets, workflows, and node cards."
        ),
    },
    "7.1": {
        "table": 12,
        "image": "fig_7_1_gaussian_controls.png",
        "width": 3.65,
        "en": (
            "Fig. 7.1. Gaussian Splat Gen on Apple MPS, showing the detected device, "
            "initializer PLY target, disabled CUDA-only True training path, and the "
            "1,000-step control."
        ),
        "zh": (
            "Fig. 7.1 Apple MPS 设备上的 Gaussian Splat Gen，展示检测到的设备、initializer "
            "PLY 目标、被禁用的 CUDA 专用 True training 路径以及 1000 步控制条。"
        ),
        "list": (
            "Fig. 7.1: Gaussian Splat Gen device, PLY target, training mode, and "
            "1,000-step controls on Apple MPS."
        ),
    },
    "8.1": {
        "table": 14,
        "image": "fig_8_1_before_after.png",
        "width": 6.2,
        "en": (
            "Fig. 8.1. Before-and-after interface comparison: the earlier version used "
            "Point Cloud Gen, Save to Library, and Reset, while the current version uses "
            "the consolidated Gaussian workflow, Save Workflow, and Clear."
        ),
        "zh": (
            "Fig. 8.1 界面前后对比：早期版本使用 Point Cloud Gen、Save to Library 和 Reset，"
            "当前版本则采用整合后的 Gaussian 工作流、Save Workflow 和 Clear。"
        ),
        "list": (
            "Fig. 8.1: Before-and-after comparison of the earlier and current workflow interfaces."
        ),
    },
}


def set_run_font(run, *, chinese: bool, italic: bool) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(8)
    run.font.italic = italic
    run._element.get_or_add_rPr().get_or_add_rFonts().set(
        qn("w:eastAsia"), "SimSun" if chinese else "Times New Roman"
    )


def replace_caption(paragraph, text: str, *, chinese: bool) -> None:
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = not chinese
    run = paragraph.add_run(text)
    set_run_font(run, chinese=chinese, italic=not chinese)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def set_cell_margins(cell, margin: int = 45) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge in ("top", "start", "bottom", "end"):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(margin))
        element.set(qn("w:type"), "dxa")


def add_alt_text(run, text: str) -> None:
    doc_properties = run._element.xpath(".//wp:docPr")
    if doc_properties:
        doc_properties[0].set("descr", text)
        doc_properties[0].set("name", text[:80])


def replace_figure_table(document: Document, spec: dict, figures_dir: Path) -> None:
    table = document.tables[spec["table"]]
    cell = table.cell(0, 0)
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, "FFFFFF")
    set_cell_margins(cell)
    row = table.rows[0]
    row.height = None
    row.height_rule = None
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(
        str(figures_dir / spec["image"]), width=Inches(spec["width"])
    )
    add_alt_text(run, spec["en"])


def main() -> None:
    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    figures_dir = Path(sys.argv[3])
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document(input_path)

    for figure_number, spec in FIGURES.items():
        replace_figure_table(document, spec, figures_dir)
        english_prefix = f"Fig. {figure_number}."
        chinese_prefix = f"Fig. {figure_number} "
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text.startswith(english_prefix) and "Placeholder:" in text:
                replace_caption(paragraph, spec["en"], chinese=False)
            elif text.startswith(chinese_prefix) and "需要展示的图片" in text:
                replace_caption(paragraph, spec["zh"], chinese=True)
            elif text.startswith(f"Fig. {figure_number}:"):
                paragraph.text = spec["list"]

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("• Replace every grey placeholder box"):
            paragraph.text = (
                "• Replace the remaining grey placeholder boxes with the required "
                "concept diagrams or experimental result images. Figures 1.1, 4.1, "
                "5.1, 6.1, 6.2, 7.1, and 8.1 already contain verified screenshots "
                "from the project website or Figma UI Kit."
            )

    document.save(output_path)


if __name__ == "__main__":
    main()
